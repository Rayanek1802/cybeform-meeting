"""
Report generation service using python-docx
"""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from datetime import datetime
from typing import Dict, Any, List
import logging

logger = logging.getLogger(__name__)


class ReportGenerator:
    """Service for generating Word reports"""
    
    def __init__(self):
        self.brand_primary = RGBColor(0x4F, 0x46, 0xE5)  # Indigo
        self.brand_secondary = RGBColor(0x7C, 0x3A, 0xED)  # Purple
        self.brand_accent = RGBColor(0x06, 0xB6, 0xD4)  # Cyan
    
    def generate_report(
        self, 
        analysis_data: Dict[str, Any], 
        transcript_segments: List[Dict[str, Any]],
        meeting_metadata: Dict[str, Any],
        output_path: str
    ) -> bool:
        """
        Generate complete Word report
        """
        try:
            doc = Document()
            
            # Ensure document is not protected or restricted
            doc.settings.odd_and_even_pages_header_footer = False
            doc.settings.different_first_page_header_footer = False
            
            # Configure document
            self._setup_document_styles(doc)
            
            # Generate content
            self._add_cover_page(doc, analysis_data, meeting_metadata)
            self._add_page_break(doc)
            self._add_table_of_contents(doc, analysis_data)
            self._add_page_break(doc)
            self._add_main_content(doc, analysis_data)
            self._add_page_break(doc)
            self._add_transcript_appendix(doc, transcript_segments)
            
            # Add headers and footers
            self._add_headers_footers(doc, meeting_metadata)
            
            # Save document
            doc.save(output_path)
            
            # Set file permissions to allow modification (666 = rw-rw-rw-)
            try:
                os.chmod(output_path, 0o666)
                logger.info(f"Report generated successfully with write permissions: {output_path}")
            except Exception as e:
                logger.warning(f"Could not set file permissions: {e}")
                logger.info(f"Report generated successfully: {output_path}")
            
            return True
            
        except Exception as e:
            logger.error(f"Error generating report: {e}")
            return False
    
    def _setup_document_styles(self, doc: Document):
        """Setup custom styles for the document"""
        styles = doc.styles
        
        # Main title style
        try:
            title_style = styles['Title']
        except KeyError:
            title_style = styles.add_style('Title', WD_STYLE_TYPE.PARAGRAPH)
        
        title_style.font.name = 'Arial'
        title_style.font.size = Pt(24)
        title_style.font.color.rgb = self.brand_primary
        title_style.font.bold = True
        
        # Heading styles
        for level in [1, 2, 3]:
            try:
                heading_style = styles[f'Heading {level}']
            except KeyError:
                heading_style = styles.add_style(f'Heading {level}', WD_STYLE_TYPE.PARAGRAPH)
            
            heading_style.font.name = 'Arial'
            heading_style.font.size = Pt(16 - level * 2)
            heading_style.font.color.rgb = self.brand_primary
            heading_style.font.bold = True
            heading_style.paragraph_format.space_before = Pt(12)
            heading_style.paragraph_format.space_after = Pt(6)
    
    def _add_cover_page(self, doc: Document, analysis_data: Dict[str, Any], meeting_metadata: Dict[str, Any]):
        """Add cover page"""
        # Logo/Brand
        title_para = doc.add_paragraph()
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_run = title_para.add_run("CybeMeeting")
        title_run.font.name = 'Arial'
        title_run.font.size = Pt(28)
        title_run.font.color.rgb = self.brand_primary
        title_run.font.bold = True
        
        # Subtitle
        subtitle_para = doc.add_paragraph()
        subtitle_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        subtitle_run = subtitle_para.add_run("Rapport de Réunion BTP")
        subtitle_run.font.name = 'Arial'
        subtitle_run.font.size = Pt(18)
        subtitle_run.font.color.rgb = self.brand_secondary
        
        # Spacing
        doc.add_paragraph("\n\n")
        
        # Meeting info
        meta = analysis_data.get("meta", {})
        info_lines = [
            f"Projet: {meta.get('projectName', 'Non spécifié')}",
            f"Type de réunion: {meta.get('meetingType', 'Autre')}",
            f"Réunion: {meta.get('meetingTitle', 'Non spécifié')}",
            f"Date: {meta.get('meetingDate', 'Non spécifié')}",
            f"Durée: {meta.get('duration', 0)} minutes",
            f"Participants détectés: {', '.join(meta.get('participantsDetected', []))}",
            f"Participants attendus: {meta.get('participantsExpected', 'Non spécifié')}"
        ]
        
        for line in info_lines:
            para = doc.add_paragraph(line)
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            para.runs[0].font.size = Pt(14)
        
        # Generation date
        doc.add_paragraph("\n\n")
        gen_para = doc.add_paragraph(f"Rapport généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}")
        gen_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        gen_para.runs[0].font.size = Pt(10)
        gen_para.runs[0].font.italic = True
    
    def _add_table_of_contents(self, doc: Document, analysis_data: Dict[str, Any]):
        """Add dynamic table of contents based on actual sections"""
        toc_para = doc.add_paragraph("SOMMAIRE", style='Heading 1')
        toc_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Always include general information
        toc_items = ["1. Informations générales"]
        
        # Generate TOC based on actual content
        if "sectionsDynamiques" in analysis_data:
            sections_data = analysis_data.get("sectionsDynamiques", {})
            
            # Define section order and French titles
            section_mappings = {
                "etatLieux": "État des lieux",
                "avancementTravaux": "Avancement des travaux", 
                "problemesIdentifies": "Problèmes identifiés",
                "decisionsStrategiques": "Décisions stratégiques",
                "objectifs": "Objectifs de la réunion",
                "actionsUrgentes": "Actions urgentes",
                "actionsReguliers": "Actions de suivi",
                "aspectsTechniques": "Aspects techniques",
                "planningEtDelais": "Planning et délais",
                "aspectsFinanciers": "Aspects financiers",
                "relationsFournisseurs": "Relations fournisseurs",
                "aspectsReglementaires": "Aspects réglementaires", 
                "communicationClient": "Communication client",
                "risquesEtMitigations": "Risques et mitigations",
                "pointsDivers": "Points divers",
                "syntheseDesAccords": "Synthèse des accords",
                "pointsEnSuspens": "Points en suspens"
            }
            
            section_counter = 2
            for section_key, title in section_mappings.items():
                if section_key in sections_data and sections_data[section_key]:
                    toc_items.append(f"{section_counter}. {title}")
                    section_counter += 1
            
            # Add chronological view if available
            if analysis_data.get("vueChronologique"):
                toc_items.append(f"{section_counter}. Vue chronologique de la réunion")
                section_counter += 1
            
            # Add analysis metrics if available
            if analysis_data.get("analysisMetrics"):
                toc_items.append(f"{section_counter}. Métriques d'analyse")
        else:
            # Fallback to legacy TOC
            toc_items.extend([
                "2. Objectifs de la réunion",
                "3. Problèmes identifiés", 
                "4. Décisions prises",
                "5. Plan d'actions",
                "6. Analyse des risques",
                "7. Points techniques BTP",
                "8. Planning et jalons",
                "9. Budget et chiffrage",
                "10. Informations diverses"
            ])
        
        # Add ANNEXE
        toc_items.append("ANNEXE - Transcription complète")
        
        for item in toc_items:
            para = doc.add_paragraph(item)
            para.paragraph_format.left_indent = Inches(0.5)
            para.paragraph_format.space_after = Pt(6)
    
    def _add_main_content(self, doc: Document, analysis_data: Dict[str, Any]):
        """Add main content sections with dynamic structure support"""
        # 1. Always start with general information
        self._add_general_information(doc, analysis_data.get("meta", {}))
        
        # Check if we have dynamic sections (new format)
        if "sectionsDynamiques" in analysis_data:
            self._add_dynamic_sections(doc, analysis_data)
        else:
            # Fallback to old format
            self._add_legacy_sections(doc, analysis_data)
        
        # Add chronological view if available
        if analysis_data.get("vueChronologique"):
            self._add_chronological_view_word(doc, analysis_data.get("vueChronologique", []))
        
        # Add analysis metrics if available
        if analysis_data.get("analysisMetrics"):
            self._add_analysis_metrics_word(doc, analysis_data.get("analysisMetrics", {}))
    
    def _add_dynamic_sections(self, doc: Document, analysis_data: Dict[str, Any]):
        """Add dynamic sections based on actual meeting content"""
        sections_data = analysis_data.get("sectionsDynamiques", {})
        
        # Define section order and French titles
        section_mappings = {
            "etatLieux": "2. État des lieux",
            "avancementTravaux": "3. Avancement des travaux", 
            "problemesIdentifies": "4. Problèmes identifiés",
            "decisionsStrategiques": "5. Décisions stratégiques",
            "objectifs": "6. Objectifs de la réunion",
            "actionsUrgentes": "7. Actions urgentes",
            "actionsReguliers": "8. Actions de suivi",
            "aspectsTechniques": "9. Aspects techniques",
            "planningEtDelais": "10. Planning et délais",
            "aspectsFinanciers": "11. Aspects financiers",
            "relationsFournisseurs": "12. Relations fournisseurs",
            "aspectsReglementaires": "13. Aspects réglementaires", 
            "communicationClient": "14. Communication client",
            "risquesEtMitigations": "15. Risques et mitigations",
            "pointsDivers": "16. Points divers",
            "syntheseDesAccords": "17. Synthèse des accords",
            "pointsEnSuspens": "18. Points en suspens"
        }
        
        section_counter = 2  # Start after general info
        
        # Process sections in logical order
        for section_key, default_title in section_mappings.items():
            if section_key in sections_data and sections_data[section_key]:
                content = sections_data[section_key]
                
                # Use numbered title
                numbered_title = f"{section_counter}. {default_title.split('. ', 1)[1]}"
                
                # Handle different content types
                if section_key in ["actionsUrgentes", "actionsReguliers"]:
                    self._add_actions_table(doc, content, numbered_title)
                elif section_key == "risquesEtMitigations":
                    self._add_risks_table(doc, content, numbered_title)
                else:
                    self._add_section(doc, numbered_title, content)
                
                section_counter += 1
        
        # Add any custom sections not in the predefined list
        for section_key, content in sections_data.items():
            if section_key not in section_mappings and content:
                # Format custom section name
                custom_title = f"{section_counter}. {section_key.replace('_', ' ').title()}"
                self._add_section(doc, custom_title, content)
                section_counter += 1
    
    def _add_legacy_sections(self, doc: Document, analysis_data: Dict[str, Any]):
        """Add sections using legacy format for backward compatibility"""
        # 2. Objectives
        self._add_section(doc, "2. Objectifs de la réunion", analysis_data.get("objectifs", []))
        
        # 3. Problems
        self._add_section(doc, "3. Problèmes identifiés", analysis_data.get("problemes", []))
        
        # 4. Decisions
        self._add_section(doc, "4. Décisions prises", analysis_data.get("decisions", []))
        
        # 5. Actions (table format)
        self._add_actions_table(doc, analysis_data.get("actions", []), "5. Plan d'actions")
        
        # 6. Risks (table format)
        self._add_risks_table(doc, analysis_data.get("risques", []), "6. Analyse des risques")
        
        # 7. Technical points
        self._add_section(doc, "7. Points techniques BTP", analysis_data.get("pointsTechniquesBTP", []))
        
        # 8. Planning
        self._add_section(doc, "8. Planning et jalons", analysis_data.get("planning", []))
        
        # 9. Budget
        self._add_section(doc, "9. Budget et chiffrage", analysis_data.get("budget_chiffrage", []))
        
        # 10. Miscellaneous
        self._add_section(doc, "10. Informations diverses", analysis_data.get("divers", []))
        
        # Exclusions (if any)
        exclusions = analysis_data.get("exclusions", [])
        if exclusions:
            self._add_section(doc, "Éléments non retenus", exclusions)
    
    def _add_chronological_view(self, doc: Document, chronological_data: List[str]):
        """Add chronological timeline of the meeting"""
        if not chronological_data:
            return
            
        doc.add_heading("Vue chronologique de la réunion", level=1)
        
        for event in chronological_data:
            if event and not event.startswith("/*"):
                para = doc.add_paragraph()
                para.style = 'List Bullet'
                para.add_run(event)
    
    def _add_analysis_metrics(self, doc: Document, metrics: Dict[str, Any]):
        """Add analysis quality metrics"""
        if not metrics:
            return
            
        doc.add_heading("Métriques d'analyse", level=1)
        
        metrics_text = [
            f"Segments analysés: {metrics.get('segmentsAnalyses', 0)}/{metrics.get('totalSegments', 0)}",
            f"Niveau de détail: {metrics.get('niveauDetaille', 'Non spécifié')}",
            f"Couverture des sujets: {metrics.get('couvertureSujets', 'Non spécifié')}",
            f"Qualité d'extraction: {metrics.get('qualiteExtraction', 'Non spécifié')}"
        ]
        
        for metric in metrics_text:
            para = doc.add_paragraph()
            para.style = 'List Bullet'
            para.add_run(metric)
    
    def _add_section(self, doc: Document, title: str, content, content_type: str = "list"):
        """Add a content section with enhanced formatting for dictionary-like content"""
        # Add title with improved styling
        heading = doc.add_heading(title, level=1)
        
        # Style the heading
        heading_format = heading.runs[0].font if heading.runs else None
        if heading_format:
            heading_format.color.rgb = RGBColor(79, 70, 229)  # Primary color
            heading_format.size = Pt(16)
            
        # Add a subtle separator line
        separator = doc.add_paragraph()
        separator_run = separator.add_run("━" * 50)
        separator_run.font.color.rgb = RGBColor(226, 232, 240)  # Light gray
        separator_run.font.size = Pt(8)
        separator.space_after = Pt(12)
        
        if content_type == "dict":
            # Handle dictionary content (meta info)
            if not content:
                doc.add_paragraph("Aucune information disponible.")
                return
            
            for key, value in content.items():
                if isinstance(value, list):
                    value = ", ".join(str(v) for v in value)
                para = doc.add_paragraph(f"{key}: {value}")
        
        elif content_type == "list":
            # Handle list content with enhanced dictionary parsing
            if not content:
                doc.add_paragraph("Aucun élément identifié.")
                return
            
            # Check if content contains dictionary-like structures
            structured_items = []
            simple_items = []
            
            for item in content:
                # Handle items that are already dict objects (from JSON parsing)
                if isinstance(item, dict):
                    structured_items.append(item)
                else:
                    item_str = str(item)
                    # Try to parse dictionary-like content from strings
                    if self._is_dict_like_string(item_str):
                        parsed_dict = self._parse_dict_string(item_str)
                        if parsed_dict:
                            structured_items.append(parsed_dict)
                        else:
                            simple_items.append(item_str)
                    else:
                        simple_items.append(item_str)
            
            # Display structured items in a table if found
            if structured_items:
                self._add_structured_table(doc, structured_items, title)
            
            # Display simple items as enhanced bullet points
            if simple_items:
                if structured_items:  # Add visual separator if we have both types
                    self._add_visual_separator(doc, "Autres éléments")
                
                for item in simple_items:
                    para = doc.add_paragraph()
                    # Add custom bullet with emoji
                    bullet_run = para.add_run("• ")
                    bullet_run.font.color.rgb = RGBColor(16, 185, 129)  # Green
                    bullet_run.font.size = Pt(14)
                    bullet_run.font.bold = True
                    
                    # Add content
                    content_run = para.add_run(str(item))
                    content_run.font.size = Pt(11)
                    content_run.font.color.rgb = RGBColor(51, 65, 85)  # Dark gray
                    
                    # Add some spacing
                    para.space_after = Pt(6)
        
        doc.add_paragraph()  # Add spacing
    
    def _add_visual_separator(self, doc: Document, text: str):
        """Add a visual separator with text"""
        separator_para = doc.add_paragraph()
        separator_para.space_before = Pt(12)
        separator_para.space_after = Pt(8)
        
        # Add text
        text_run = separator_para.add_run(text)
        text_run.font.bold = True
        text_run.font.color.rgb = RGBColor(124, 58, 237)  # Secondary color
        text_run.font.size = Pt(12)
    
    def _is_dict_like_string(self, text: str) -> bool:
        """Check if a string contains dictionary-like content"""
        import re
        # Look for patterns like {'key': 'value'} or key: value
        dict_patterns = [
            r"\{['\"][\w\s]+['\"]:\s*['\"].*?['\"][,}]",  # {'key': 'value'}
            r"['\"][\w\s]+['\"]:\s*['\"].*?['\"]",        # 'key': 'value'
            r"[\w\s]+:\s*['\"].*?['\"]"                   # key: 'value'
        ]
        for pattern in dict_patterns:
            if re.search(pattern, text):
                return True
        return False
    
    def _parse_dict_string(self, text: str) -> dict:
        """Parse dictionary-like string into structured data"""
        import re
        parsed_dict = {}
        
        # Enhanced patterns to capture complete content including apostrophes and long text
        patterns = [
            # Pattern for complete dictionary with proper quote handling
            r"['\"](\w+)['\"]:\s*['\"]([^'\"]*(?:[^'\"]+[^'\"]*)*)['\"]",  # Basic pattern
            # Pattern that handles escaped quotes and longer content
            r"['\"]([^'\"]+)['\"]:\s*['\"]([^'\"]*(?:\\.[^'\"]*)*)['\"]",   # With escaped chars
            # Pattern for dictionary values that might contain apostrophes
            r"['\"](\w+)['\"]:\s*['\"]([^'\"]*?)['\"](?:\s*[,}])",         # Non-greedy with delimiters
        ]
        
        # First try to extract the full dictionary structure
        dict_match = re.search(r'\{([^{}]+)\}', text)
        if dict_match:
            dict_content = dict_match.group(1)
            
            # Split by comma but be careful of commas inside quotes
            items = []
            current_item = ""
            in_quotes = False
            quote_char = None
            
            i = 0
            while i < len(dict_content):
                char = dict_content[i]
                
                # Check for escaped characters
                if char == '\\' and i + 1 < len(dict_content):
                    # Add the escape sequence as-is, we'll handle it later
                    current_item += char + dict_content[i + 1]
                    i += 2
                    continue
                    
                if char in ["'", '"'] and not in_quotes:
                    in_quotes = True
                    quote_char = char
                elif char == quote_char and in_quotes:
                    in_quotes = False
                    quote_char = None
                elif char == ',' and not in_quotes:
                    if current_item.strip():
                        items.append(current_item.strip())
                    current_item = ""
                    i += 1
                    continue
                
                current_item += char
                i += 1
            
            if current_item.strip():
                items.append(current_item.strip())
            
            # Parse each item
            for item in items:
                # Extract key-value from each item, handling escaped quotes
                # This pattern handles escaped quotes properly by using atomic groups and backtracking
                kv_match = re.search(r"['\"]([^'\"]+)['\"]:\s*['\"](((?:[^'\"\\]|\\.)*))['\"]", item)
                if kv_match:
                    key, value = kv_match.groups()[:2]  # Take only key and value
                    # Unescape common escape sequences
                    value = value.replace("\\'", "'").replace('\\"', '"').replace("\\\\", "\\")
                    parsed_dict[key] = value
        
        # Fallback to simple patterns if dictionary parsing failed
        if not parsed_dict:
            for pattern in patterns:
                matches = re.findall(pattern, text)
                for key, value in matches:
                    parsed_dict[key] = value
        
        return parsed_dict
    
    def _add_structured_table(self, doc: Document, structured_items: list, title: str):
        """Add a table for structured dictionary-like items"""
        if not structured_items:
            return
        
        # Determine table columns based on common keys
        all_keys = set()
        for item in structured_items:
            all_keys.update(item.keys())
        
        # Common keys mapping to French
        key_mappings = {
            'decision': 'Décision',
            'detail': 'Détail',
            'risk': 'Risque',
            'recommendation': 'Recommandation',
            'point': 'Point',
            'participant': 'Participant',
            'context': 'Contexte',
            'contexte': 'Contexte',
            'contexteTemporel': 'Temps',
            'priority': 'Priorité',
            'priorite': 'Priorité',
            'action': 'Action',
            'tache': 'Tâche',
            'responsable': 'Responsable',
            'echeance': 'Échéance'
        }
        
        # Filter and order columns
        ordered_keys = []
        main_keys = ['decision', 'detail', 'risk', 'recommendation', 'point', 'participant', 'action', 'tache']
        
        # Add main content key first
        for key in main_keys:
            if key in all_keys:
                ordered_keys.append(key)
                break
        
        # Add other important keys
        other_important_keys = ['context', 'contexte', 'contexteTemporel', 'priority', 'priorite', 'responsable', 'echeance']
        for key in other_important_keys:
            if key in all_keys and key not in ordered_keys:
                ordered_keys.append(key)
        
        # Add remaining keys
        for key in sorted(all_keys):
            if key not in ordered_keys:
                ordered_keys.append(key)
        
        # Create enhanced table with professional styling
        num_cols = min(len(ordered_keys), 4)  # Limit to 4 columns for readability
        table = doc.add_table(rows=1, cols=num_cols)
        
        # Apply modern table style
        table.style = 'Light Grid Accent 1'
        table.allow_autofit = False  # Disable autofit to allow custom column widths
        
        # Enhanced column width distribution for better readability
        from docx.shared import Inches
        total_width = 7.5  # Total page width minus margins
        if num_cols == 1:
            table.columns[0].width = Inches(total_width)
        elif num_cols == 2:
            table.columns[0].width = Inches(total_width * 0.7)  # Main content
            table.columns[1].width = Inches(total_width * 0.3)  # Context/Time
        elif num_cols == 3:
            table.columns[0].width = Inches(total_width * 0.5)  # Main content
            table.columns[1].width = Inches(total_width * 0.3)  # Context
            table.columns[2].width = Inches(total_width * 0.2)  # Time
        elif num_cols >= 4:
            table.columns[0].width = Inches(total_width * 0.4)  # Main content
            table.columns[1].width = Inches(total_width * 0.25)  # Context
            table.columns[2].width = Inches(total_width * 0.2)   # Time
            table.columns[3].width = Inches(total_width * 0.15)  # Additional
        
        # Enhanced header row with modern styling
        header_cells = table.rows[0].cells
        for i, key in enumerate(ordered_keys[:num_cols]):
            # Set header text with icon
            header_text = key_mappings.get(key, key.title())
            
            # Set header text without icons for professional look
            header_cells[i].text = header_text
            
            # Enhanced header formatting
            header_para = header_cells[i].paragraphs[0]
            header_run = header_para.runs[0]
            header_run.font.bold = True
            header_run.font.color.rgb = RGBColor(255, 255, 255)  # White text
            header_run.font.size = Pt(11)
            header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Professional header background
            try:
                from docx.oxml.shared import qn
                from docx.oxml.parser import parse_xml
                from docx.oxml.ns import nsdecls
                
                tcPr = header_cells[i]._tc.get_or_add_tcPr()
                # Gradient-like effect with brand primary color
                shd = parse_xml(r'<w:shd {} w:fill="4F46E5"/>'.format(nsdecls('w')))
                tcPr.append(shd)
                
                # Add padding to header cells
                tcMar = tcPr.find(qn('w:tcMar'))
                if tcMar is None:
                    tcMar = parse_xml(r'<w:tcMar {}><w:top w:w="120" w:type="dxa"/><w:left w:w="120" w:type="dxa"/><w:bottom w:w="120" w:type="dxa"/><w:right w:w="120" w:type="dxa"/></w:tcMar>'.format(nsdecls('w')))
                    tcPr.append(tcMar)
            except:
                # Fallback styling
                header_run.font.color.rgb = self.brand_primary
        
        # Enhanced data rows with modern styling
        for row_idx, item in enumerate(structured_items):
            row_cells = table.add_row().cells
            
            # Add subtle alternating row colors
            is_even_row = row_idx % 2 == 0
            
            for i, key in enumerate(ordered_keys[:num_cols]):
                value = item.get(key, '')
                
                # Format temporal context with improved parsing
                if key == 'contexteTemporel' and value:
                    import re
                    # Handle multiple time formats
                    time_patterns = [
                        r'\[(\d{2}:\d{2})-(\d{2}:\d{2})\]',  # [00:00-16:07]
                        r'(\d{2}:\d{2})\s*-\s*(\d{2}:\d{2})',  # 00:00 - 16:07
                        r'(\d{2}:\d{2})',  # 00:00
                    ]
                    
                    for pattern in time_patterns:
                        time_match = re.search(pattern, str(value))
                        if time_match:
                            if len(time_match.groups()) == 2:
                                value = f"{time_match.group(1)} - {time_match.group(2)}"
                            else:
                                value = time_match.group(1)
                            break
                
                # Enhanced cell content formatting
                cell_para = row_cells[i].paragraphs[0]
                cell_para.clear()  # Clear existing content
                
                # Add formatted content
                if value:
                    cell_run = cell_para.add_run(str(value))
                    cell_run.font.size = Pt(10)
                    
                    # Style based on column type
                    if i == 0:  # Main content column
                        cell_run.font.bold = True
                        cell_run.font.color.rgb = RGBColor(31, 41, 55)  # Dark gray
                        cell_run.font.size = Pt(11)
                    elif key == 'context' or key == 'contexte':  # Context column
                        cell_run.font.italic = True
                        cell_run.font.color.rgb = RGBColor(107, 114, 128)  # Medium gray
                    elif key == 'contexteTemporel':  # Time column
                        cell_run.font.color.rgb = RGBColor(79, 70, 229)  # Brand color
                        cell_run.font.size = Pt(9)
                    else:
                        cell_run.font.color.rgb = RGBColor(75, 85, 99)  # Gray
                
                # Enhanced cell alignment and spacing
                cell_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                cell_para.space_before = Pt(4)
                cell_para.space_after = Pt(4)
                
                # Advanced cell properties for better presentation
                try:
                    from docx.oxml.shared import qn
                    from docx.oxml import OxmlElement
                    from docx.oxml.parser import parse_xml
                    from docx.oxml.ns import nsdecls
                    
                    tcPr = row_cells[i]._tc.get_or_add_tcPr()
                    
                    # Enable proper text wrapping
                    tcW = tcPr.find(qn('w:tcW'))
                    if tcW is None:
                        tcW = OxmlElement(qn('w:tcW'))
                        tcW.set(qn('w:type'), 'auto')
                        tcPr.append(tcW)
                    
                    # Add subtle alternating row backgrounds
                    if is_even_row:
                        shd = tcPr.find(qn('w:shd'))
                        if shd is None:
                            shd = parse_xml(r'<w:shd {} w:fill="F8FAFC"/>'.format(nsdecls('w')))
                            tcPr.append(shd)
                    
                    # Add cell padding for better readability
                    tcMar = tcPr.find(qn('w:tcMar'))
                    if tcMar is None:
                        tcMar = parse_xml(r'<w:tcMar {}><w:top w:w="100" w:type="dxa"/><w:left w:w="100" w:type="dxa"/><w:bottom w:w="100" w:type="dxa"/><w:right w:w="100" w:type="dxa"/></w:tcMar>'.format(nsdecls('w')))
                        tcPr.append(tcMar)
                        
                    # Set vertical alignment to top for better multi-line content
                    vAlign = tcPr.find(qn('w:vAlign'))
                    if vAlign is None:
                        vAlign = OxmlElement(qn('w:vAlign'))
                        vAlign.set(qn('w:val'), 'top')
                        tcPr.append(vAlign)
                        
                except Exception:
                    pass  # Continue if XML manipulation fails
        
    
    
    def _add_actions_table(self, doc: Document, actions: List[Dict[str, Any]], title: str = "5. Plan d'actions"):
        """Add enhanced actions table with full information"""
        doc.add_heading(title, level=1)
        
        if not actions:
            doc.add_paragraph("Aucune action définie.")
            return
        
        # Create table with 5 columns including Context and Priority
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        table.allow_autofit = True
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Action'
        header_cells[1].text = 'Responsable'
        header_cells[2].text = 'Échéance'
        header_cells[3].text = 'Priorité'
        header_cells[4].text = 'Contexte'
        
        # Style header
        for cell in header_cells:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = self.brand_primary
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            # Add light background to header
            try:
                from docx.oxml.shared import qn
                tcPr = cell._tc.get_or_add_tcPr()
                shd = tcPr.find(qn('w:shd'))
                if shd is None:
                    from docx.oxml.parser import parse_xml
                    from docx.oxml.ns import nsdecls
                    shd = parse_xml(r'<w:shd {} w:fill="F1F5F9"/>'.format(nsdecls('w')))
                    tcPr.append(shd)
            except:
                pass  # Skip background color if it fails
        
        # Add action rows
        for action in actions:
            row_cells = table.add_row().cells
            
            # Action description
            action_text = action.get('action', action.get('tache', ''))
            row_cells[0].text = str(action_text)
            row_cells[0].paragraphs[0].runs[0].font.size = Pt(10)
            
            # Responsible person
            row_cells[1].text = str(action.get('responsable', 'Non assigné'))
            
            # Due date
            echeance = action.get('echeance', 'Non définie')
            row_cells[2].text = str(echeance) if echeance else 'Non définie'
            
            # Priority
            priorite = action.get('priorite', 'Moyenne')
            row_cells[3].text = str(priorite)
            
            # Priority color coding
            if str(priorite).lower() in ['haute', 'urgent', 'urgente']:
                row_cells[3].paragraphs[0].runs[0].font.color.rgb = RGBColor(239, 68, 68)  # Red
                row_cells[3].paragraphs[0].runs[0].font.bold = True
            elif str(priorite).lower() in ['moyenne', 'normal', 'normale']:
                row_cells[3].paragraphs[0].runs[0].font.color.rgb = RGBColor(245, 158, 11)  # Orange
            else:
                row_cells[3].paragraphs[0].runs[0].font.color.rgb = RGBColor(16, 185, 129)  # Green
            
            # Context
            contexte = action.get('contexte', action.get('dependances', ''))
            row_cells[4].text = str(contexte) if contexte else 'Non spécifié'
            row_cells[4].paragraphs[0].runs[0].font.size = Pt(9)
        
        doc.add_paragraph()  # Add spacing
    
    def _add_risks_table(self, doc: Document, risks: List[Dict[str, Any]], title: str = "6. Analyse des risques"):
        """Add enhanced risks table with full information"""
        doc.add_heading(title, level=1)
        
        if not risks:
            doc.add_paragraph("Aucun risque identifié.")
            return
        
        # Create table with 6 columns for complete risk information
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        table.allow_autofit = True
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Risque'
        header_cells[1].text = 'Catégorie'
        header_cells[2].text = 'Probabilité'
        header_cells[3].text = 'Impact'
        header_cells[4].text = 'Mitigation'
        header_cells[5].text = 'Responsable'
        
        # Style header
        for cell in header_cells:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = self.brand_primary
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            # Add light background to header
            try:
                from docx.oxml.shared import qn
                tcPr = cell._tc.get_or_add_tcPr()
                shd = tcPr.find(qn('w:shd'))
                if shd is None:
                    from docx.oxml.parser import parse_xml
                    from docx.oxml.ns import nsdecls
                    shd = parse_xml(r'<w:shd {} w:fill="F1F5F9"/>'.format(nsdecls('w')))
                    tcPr.append(shd)
            except:
                pass  # Skip background color if it fails
        
        # Add risk rows
        for risk in risks:
            row_cells = table.add_row().cells
            
            # Risk description
            row_cells[0].text = str(risk.get('risque', 'Non spécifié'))
            row_cells[0].paragraphs[0].runs[0].font.size = Pt(10)
            
            # Category
            row_cells[1].text = str(risk.get('categorie', 'Général'))
            
            # Probability with color coding
            probabilite = risk.get('probabilite', 'Moyenne')
            row_cells[2].text = str(probabilite)
            
            # Color code probability
            if str(probabilite).lower() in ['élevée', 'haute', 'high']:
                row_cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(239, 68, 68)  # Red
                row_cells[2].paragraphs[0].runs[0].font.bold = True
            elif str(probabilite).lower() in ['moyenne', 'medium']:
                row_cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(245, 158, 11)  # Orange
            else:
                row_cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(16, 185, 129)  # Green
            
            # Impact
            row_cells[3].text = str(risk.get('impact', 'Non spécifié'))
            row_cells[3].paragraphs[0].runs[0].font.size = Pt(9)
            
            # Mitigation
            mitigation = risk.get('mitigations', risk.get('mitigation', 'Non spécifié'))
            row_cells[4].text = str(mitigation)
            row_cells[4].paragraphs[0].runs[0].font.size = Pt(9)
            
            # Responsible
            responsable = risk.get('responsableRisque', risk.get('responsable', 'Non assigné'))
            row_cells[5].text = str(responsable)
        
        doc.add_paragraph()  # Add spacing
    
    def _add_transcript_appendix(self, doc: Document, transcript_segments: List[Dict[str, Any]]):
        """Add transcript appendix"""
        doc.add_heading("ANNEXE - Transcription complète", level=1)
        
        if not transcript_segments:
            doc.add_paragraph("Transcription non disponible.")
            return
        
        current_speaker = None
        
        for segment in transcript_segments:
            speaker = segment.get("speaker", "INCONNU")
            text = segment.get("text", "").strip()
            start_time = segment.get("start_time", 0)
            
            if not text:
                continue
            
            # Add speaker header if changed
            if speaker != current_speaker:
                if current_speaker is not None:
                    doc.add_paragraph()  # Add spacing between speakers
                
                speaker_para = doc.add_paragraph(f"{speaker}")
                speaker_para.runs[0].font.bold = True
                speaker_para.runs[0].font.color.rgb = self.brand_primary
                current_speaker = speaker
            
            # Add timestamp and text
            minutes = int(start_time // 60)
            seconds = int(start_time % 60)
            timestamp = f"[{minutes:02d}:{seconds:02d}]"
            
            para = doc.add_paragraph(f"{timestamp} {text}")
            para.paragraph_format.left_indent = Inches(0.5)
    
    def _add_headers_footers(self, doc: Document, meeting_metadata: Dict[str, Any]):
        """Add headers and footers"""
        # Header
        header = doc.sections[0].header
        header_para = header.paragraphs[0]
        header_para.text = "CybeMeeting - Rapport de Réunion BTP"
        header_para.runs[0].font.size = Pt(10)
        header_para.runs[0].font.color.rgb = self.brand_primary
        
        # Footer
        footer = doc.sections[0].footer
        footer_para = footer.paragraphs[0]
        project_name = meeting_metadata.get("project_name", "Projet BTP")
        footer_para.text = f"{project_name} - {datetime.now().strftime('%d/%m/%Y')}"
        footer_para.runs[0].font.size = Pt(9)
        footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    def _add_page_break(self, doc: Document):
        """Add page break"""
        doc.add_page_break()
    
    def generate_html_preview(self, analysis_data: Dict[str, Any], meeting_metadata: Dict[str, Any]) -> str:
        """Generate HTML preview of the report"""
        try:
            meta = analysis_data.get("meta", {})
            
            html = f"""
            <!DOCTYPE html>
            <html lang="fr">
            <head>
                <meta charset="UTF-8">
                <meta name="viewport" content="width=device-width, initial-scale=1.0">
                <title>Rapport de Réunion - {meta.get('meetingTitle', 'Réunion')}</title>
                <style>
                    /* Modern CSS Variables */
                    :root {{
                        --primary-color: #4F46E5;
                        --secondary-color: #7C3AED;
                        --accent-color: #06B6D4;
                        --success-color: #10B981;
                        --warning-color: #F59E0B;
                        --error-color: #EF4444;
                        --gray-50: #F8FAFC;
                        --gray-100: #F1F5F9;
                        --gray-200: #E2E8F0;
                        --gray-600: #475569;
                        --gray-700: #334155;
                        --shadow-sm: 0 1px 2px 0 rgba(0, 0, 0, 0.05);
                        --shadow-md: 0 4px 6px -1px rgba(0, 0, 0, 0.1), 0 2px 4px -1px rgba(0, 0, 0, 0.06);
                        --shadow-lg: 0 10px 15px -3px rgba(0, 0, 0, 0.1), 0 4px 6px -2px rgba(0, 0, 0, 0.05);
                        --border-radius: 0.75rem;
                    }}

                    /* Base Styles */
                    * {{ box-sizing: border-box; }}
                    body {{ 
                        font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', 'Roboto', 'Helvetica Neue', Arial, sans-serif;
                        line-height: 1.7;
                        margin: 0;
                        padding: 2rem;
                        background: linear-gradient(135deg, var(--gray-50) 0%, #e0e7ff 100%);
                        color: var(--gray-700);
                        font-size: 15px;
                    }}

                    /* Main Container */
                    .container {{ 
                        max-width: 900px;
                        margin: 0 auto;
                        background: white;
                        padding: 2.5rem;
                        border-radius: var(--border-radius);
                        box-shadow: var(--shadow-lg);
                        border: 1px solid var(--gray-200);
                    }}

                    /* Header */
                    .header {{ 
                        text-align: center;
                        margin-bottom: 2.5rem;
                        padding-bottom: 1.5rem;
                        position: relative;
                    }}
                    .header::after {{
                        content: '';
                        position: absolute;
                        bottom: 0;
                        left: 50%;
                        transform: translateX(-50%);
                        width: 100px;
                        height: 4px;
                        background: linear-gradient(90deg, var(--primary-color), var(--secondary-color));
                        border-radius: 2px;
                    }}
                    .brand {{ 
                        color: var(--primary-color);
                        font-size: 2rem;
                        font-weight: 800;
                        margin-bottom: 0.5rem;
                        letter-spacing: -0.025em;
                    }}
                    .subtitle {{ 
                        color: var(--secondary-color);
                        font-size: 1.25rem;
                        font-weight: 500;
                        margin-bottom: 1rem;
                    }}

                    /* Meta Information */
                    .meta-info {{ 
                        background: linear-gradient(135deg, var(--gray-50), var(--gray-100));
                        padding: 1.5rem;
                        border-radius: calc(var(--border-radius) - 0.25rem);
                        margin-bottom: 2rem;
                        border: 1px solid var(--gray-200);
                        display: grid;
                        grid-template-columns: repeat(auto-fit, minmax(250px, 1fr));
                        gap: 1rem;
                    }}
                    .meta-info div {{ 
                        display: flex;
                        align-items: center;
                        gap: 0.5rem;
                        color: var(--gray-600);
                        font-weight: 500;
                    }}

                    /* Typography */
                    h1 {{ 
                        color: var(--primary-color);
                        font-size: 1.875rem;
                        font-weight: 700;
                        margin: 2rem 0 1.5rem 0;
                        padding-bottom: 0.75rem;
                        border-bottom: 2px solid var(--gray-200);
                        letter-spacing: -0.025em;
                    }}
                    h2 {{ 
                        color: var(--secondary-color);
                        font-size: 1.5rem;
                        font-weight: 600;
                        margin: 2rem 0 1rem 0;
                        display: flex;
                        align-items: center;
                        gap: 0.75rem;
                        letter-spacing: -0.025em;
                    }}
                    h2::before {{
                        content: '';
                        width: 4px;
                        height: 2rem;
                        background: linear-gradient(180deg, var(--secondary-color), var(--accent-color));
                        border-radius: 2px;
                    }}

                    /* Sections */
                    .section {{ 
                        margin-bottom: 2rem;
                        background: white;
                        border-radius: calc(var(--border-radius) - 0.25rem);
                        overflow: hidden;
                    }}
                    .section-content {{ 
                        background: var(--gray-50);
                        padding: 1.5rem;
                        border-radius: calc(var(--border-radius) - 0.25rem);
                        border: 1px solid var(--gray-200);
                    }}

                    /* Lists */
                    ul {{ 
                        padding-left: 0;
                        list-style: none;
                    }}
                    li {{ 
                        margin-bottom: 1rem;
                        padding: 1rem 1.25rem;
                        background: white;
                        border-radius: calc(var(--border-radius) - 0.5rem);
                        border: 1px solid var(--gray-200);
                        box-shadow: var(--shadow-sm);
                        transition: all 0.2s ease;
                        position: relative;
                        padding-left: 2.5rem;
                    }}
                    li::before {{
                        content: '✓';
                        position: absolute;
                        left: 1rem;
                        top: 50%;
                        transform: translateY(-50%);
                        color: var(--success-color);
                        font-weight: bold;
                        font-size: 1rem;
                    }}
                    li:hover {{
                        transform: translateY(-1px);
                        box-shadow: var(--shadow-md);
                        border-color: var(--primary-color);
                    }}

                    /* Tables */
                    table {{ 
                        width: 100%;
                        border-collapse: separate;
                        border-spacing: 0;
                        margin: 1.5rem 0;
                        border-radius: calc(var(--border-radius) - 0.25rem);
                        overflow: hidden;
                        box-shadow: var(--shadow-sm);
                        border: 1px solid var(--gray-200);
                    }}
                    th {{ 
                        background: linear-gradient(135deg, var(--primary-color), var(--secondary-color));
                        color: white;
                        padding: 1rem;
                        text-align: left;
                        font-weight: 600;
                        font-size: 0.875rem;
                        text-transform: uppercase;
                        letter-spacing: 0.05em;
                    }}
                    td {{ 
                        border: none;
                        border-bottom: 1px solid var(--gray-200);
                        padding: 1rem;
                        vertical-align: top;
                        background: white;
                    }}
                    tr:nth-child(even) td {{
                        background: var(--gray-50);
                    }}
                    tr:hover td {{
                        background: #e0f2fe;
                    }}

                    /* Special Styles */
                    .empty-section {{ 
                        color: var(--gray-600);
                        font-style: italic;
                        text-align: center;
                        padding: 2rem;
                        background: var(--gray-50);
                        border-radius: calc(var(--border-radius) - 0.5rem);
                    }}

                    /* Strong and Emphasis */
                    strong {{
                        color: var(--gray-700);
                        font-weight: 700;
                    }}
                    em {{
                        color: var(--gray-600);
                        font-style: italic;
                    }}

                    /* Responsive Design */
                    @media (max-width: 768px) {{
                        body {{ padding: 1rem; }}
                        .container {{ padding: 1.5rem; }}
                        .brand {{ font-size: 1.5rem; }}
                        .subtitle {{ font-size: 1.125rem; }}
                        h1 {{ font-size: 1.5rem; }}
                        h2 {{ font-size: 1.25rem; }}
                        .meta-info {{ grid-template-columns: 1fr; }}
                    }}
                </style>
            </head>
            <body>
                <div class="container">
                    <div class="header">
                        <div class="brand">CybeMeeting</div>
                        <div class="subtitle">Rapport de Réunion BTP</div>
                    </div>
                    
                    <div class="meta-info">
                        <div><strong>Projet:</strong> {meta.get('projectName', 'Non spécifié')}</div>
                        <div><strong>Réunion:</strong> {meta.get('meetingTitle', 'Non spécifié')}</div>
                        <div><strong>Date:</strong> {meta.get('meetingDate', 'Non spécifié')}</div>
                        <div><strong>Durée:</strong> {meta.get('duration', 0)} minutes</div>
                        <div><strong>Participants:</strong> {', '.join(meta.get('participantsDetected', []))}</div>
                    </div>
            """
            
            # Add user instructions if present
            user_instructions = meta.get('userInstructions', '')
            if user_instructions:
                html += f"""
                    <div class="user-instructions" style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; padding: 15px; border-radius: 8px; margin-bottom: 25px;">
                        <strong style="color: #f1f5f9;">Instructions spécifiques:</strong> {user_instructions}
                    </div>
                """
            
            # Check if we have dynamic sections (new format)
            if "sectionsDynamiques" in analysis_data:
                html += self._generate_dynamic_sections_html(analysis_data.get("sectionsDynamiques", {}))
            else:
                # Fallback to old format
                html += self._generate_legacy_sections_html(analysis_data)
            
            # Add chronological view if available
            if analysis_data.get("vueChronologique"):
                html += self._generate_chronological_html(analysis_data.get("vueChronologique", []))
            
            # Add analysis metrics if available
            if analysis_data.get("analysisMetrics"):
                html += self._generate_metrics_html(analysis_data.get("analysisMetrics", {}))
            
            html += """
                    <div style="text-align: center; margin-top: 30px; padding-top: 20px; border-top: 1px solid #e2e8f0; color: #64748b; font-size: 12px;">
                        Rapport généré par CybeMeeting le """ + datetime.now().strftime('%d/%m/%Y à %H:%M') + """
                    </div>
                </div>
            </body>
            </html>
            """
            
            return html
            
        except Exception as e:
            logger.error(f"Error generating HTML preview: {e}")
            return "<html><body><h1>Erreur lors de la génération de l'aperçu</h1></body></html>"
    
    def _generate_dynamic_sections_html(self, sections_data: Dict[str, Any]) -> str:
        """Generate HTML for dynamic sections"""
        html = ""
        
        # Define section order and French titles
        section_mappings = {
            "etatLieux": "État des lieux",
            "avancementTravaux": "Avancement des travaux", 
            "problemesIdentifies": "Problèmes identifiés",
            "decisionsStrategiques": "Décisions stratégiques",
            "objectifs": "Objectifs de la réunion",
            "actionsUrgentes": "Actions urgentes",
            "actionsReguliers": "Actions de suivi",
            "aspectsTechniques": "Aspects techniques",
            "planningEtDelais": "Planning et délais",
            "aspectsFinanciers": "Aspects financiers",
            "relationsFournisseurs": "Relations fournisseurs",
            "aspectsReglementaires": "Aspects réglementaires", 
            "communicationClient": "Communication client",
            "risquesEtMitigations": "Risques et mitigations",
            "pointsDivers": "Points divers",
            "syntheseDesAccords": "Synthèse des accords",
            "pointsEnSuspens": "Points en suspens"
        }
        
        # Process sections in logical order
        for section_key, title in section_mappings.items():
            if section_key in sections_data and sections_data[section_key]:
                content = sections_data[section_key]
                html += f'<div class="section"><h2 style="color: #7C3AED; margin-top: 25px; border-left: 4px solid #7C3AED; padding-left: 15px;">{title}</h2>'
                html += '<div class="section-content" style="background: #fafafa; padding: 15px; border-radius: 6px;">'
                
                # Handle different content types
                if section_key in ["actionsUrgentes", "actionsReguliers"]:
                    html += self._generate_actions_table_html(content, section_key == "actionsUrgentes")
                elif section_key == "risquesEtMitigations":
                    html += self._generate_risks_table_html(content)
                else:
                    html += self._generate_list_html(content)
                
                html += '</div></div>'
        
        # Add any custom sections not in the predefined list
        for section_key, content in sections_data.items():
            if section_key not in section_mappings and content:
                custom_title = self._format_section_title(section_key)
                html += f'<div class="section"><h2 style="color: #7C3AED; margin-top: 25px; border-left: 4px solid #7C3AED; padding-left: 15px;">{custom_title}</h2>'
                html += '<div class="section-content" style="background: #fafafa; padding: 15px; border-radius: 6px;">'
                html += self._generate_list_html(content)
                html += '</div></div>'
        
        return html
    
    def _format_section_title(self, section_key: str) -> str:
        """Format section key to readable title (handles camelCase and snake_case)"""
        import re
        
        # First replace underscores with spaces
        title = section_key.replace('_', ' ')
        
        # Then split camelCase - insert space before uppercase letters
        title = re.sub(r'([a-z])([A-Z])', r'\1 \2', title)
        
        # Convert to title case
        title = title.title()
        
        return title
    
    def _generate_legacy_sections_html(self, analysis_data: Dict[str, Any]) -> str:
        """Generate HTML for legacy format"""
        html = ""
        sections = [
            ("Objectifs de la réunion", analysis_data.get("objectifs", []), "list"),
            ("Problèmes identifiés", analysis_data.get("problemes", []), "list"),
            ("Décisions prises", analysis_data.get("decisions", []), "list"),
            ("Plan d'actions", analysis_data.get("actions", []), "actions_table"),
            ("Analyse des risques", analysis_data.get("risques", []), "risks_table"),
            ("Points techniques BTP", analysis_data.get("pointsTechniquesBTP", []), "list"),
            ("Planning et jalons", analysis_data.get("planning", []), "list"),
            ("Budget et chiffrage", analysis_data.get("budget_chiffrage", []), "list"),
            ("Informations diverses", analysis_data.get("divers", []), "list")
        ]
        
        for title, content, content_type in sections:
            html += f'<div class="section"><h2 style="color: #7C3AED; margin-top: 25px;">{title}</h2>'
            html += '<div class="section-content" style="background: #fafafa; padding: 15px; border-radius: 6px;">'
            
            if not content:
                html += '<div style="color: #64748b; font-style: italic; text-align: center; padding: 20px;">Aucun élément identifié.</div>'
            elif content_type == "list":
                html += self._generate_list_html(content)
            elif content_type == "actions_table":
                html += self._generate_actions_table_html(content, False)
            elif content_type == "risks_table":
                html += self._generate_risks_table_html(content)
            
            html += '</div></div>'
        
        return html
    
    def _generate_list_html(self, items: list) -> str:
        """Generate HTML for list items"""
        if not items:
            return '<div style="color: #64748b; font-style: italic; text-align: center; padding: 20px;">Aucun élément identifié.</div>'
        
        html = "<ul style='padding-left: 20px;'>"
        for item in items:
            if item and not str(item).startswith("/*"):
                # Format dictionary objects properly
                if isinstance(item, dict):
                    formatted_item = self._format_dict_for_html(item)
                    html += f"<li style='margin-bottom: 8px; line-height: 1.4;'>{formatted_item}</li>"
                else:
                    html += f"<li style='margin-bottom: 8px; line-height: 1.4;'>{item}</li>"
        html += "</ul>"
        return html
    
    def _format_dict_for_html(self, item_dict: dict) -> str:
        """Format a dictionary object for HTML display"""
        # Map of common field keys to display names
        field_mapping = {
            'decision': '🎯 Décision',
            'detail': '🔧 Détail technique',
            'risk': '⚠️ Risque',
            'recommendation': '💡 Recommandation',
            'point': '📌 Point important',
            'action': '✅ Action',
            'tache': '✅ Tâche',
            'context': '📝 Contexte',
            'contexte': '📝 Contexte',
            'contexteTemporel': '⏰ Moment',
            'responsable': '👤 Responsable',
            'echeance': '📅 Échéance',
            'priorite': '🔥 Priorité'
        }
        
        # Extract main content (first field that's not context/time)
        main_content = None
        context_info = []
        
        for key, value in item_dict.items():
            if key in ['decision', 'detail', 'risk', 'recommendation', 'point', 'action', 'tache']:
                main_content = f"<strong>{value}</strong>"
            elif key in ['context', 'contexte']:
                if value:
                    context_info.append(f"<em>Contexte: {value}</em>")
            elif key == 'contexteTemporel':
                if value:
                    # Format time context nicely
                    time_str = value.replace('[', '').replace(']', '')
                    context_info.append(f"<span style='color: #6B7280; font-size: 0.9em;'>⏱️ {time_str}</span>")
            elif value:  # Other fields
                field_name = field_mapping.get(key, key.capitalize())
                context_info.append(f"{field_name}: {value}")
        
        # Build formatted string
        if main_content:
            result = main_content
            if context_info:
                result += f"<div style='margin-top: 4px; padding-left: 10px; color: #6B7280; font-size: 0.95em;'>{' • '.join(context_info)}</div>"
            return result
        else:
            # Fallback: just show all fields nicely
            formatted_fields = []
            for key, value in item_dict.items():
                if value:
                    field_name = field_mapping.get(key, key.capitalize())
                    formatted_fields.append(f"{field_name}: {value}")
            return "<br>".join(formatted_fields)
    
    def _generate_actions_table_html(self, actions: list, is_urgent: bool = False) -> str:
        """Generate HTML table for actions"""
        if not actions:
            return '<div style="color: #64748b; font-style: italic; text-align: center; padding: 20px;">Aucune action définie.</div>'
        
        html = """
        <table style="width: 100%; border-collapse: collapse; margin: 15px 0;">
            <thead>
                <tr><th style="border: 1px solid #e2e8f0; padding: 12px; text-align: left; background: #f8fafc; font-weight: bold; color: #4F46E5;">Action</th><th style="border: 1px solid #e2e8f0; padding: 12px; text-align: left; background: #f8fafc; font-weight: bold; color: #4F46E5;">Responsable</th><th style="border: 1px solid #e2e8f0; padding: 12px; text-align: left; background: #f8fafc; font-weight: bold; color: #4F46E5;">Échéance</th><th style="border: 1px solid #e2e8f0; padding: 12px; text-align: left; background: #f8fafc; font-weight: bold; color: #4F46E5;">Contexte</th></tr>
            </thead>
            <tbody>
        """
        
        for action in actions:
            border_color = "#ef4444" if is_urgent else "#f59e0b"
            bg_color = "#fef2f2" if is_urgent else "#fefbf2"
            html += f"""
            <tr style="border-left: 4px solid {border_color}; background: {bg_color};">
                <td style="border: 1px solid #e2e8f0; padding: 12px;">{action.get('action', action.get('tache', ''))}</td>
                <td style="border: 1px solid #e2e8f0; padding: 12px;">{action.get('responsable', '')}</td>
                <td style="border: 1px solid #e2e8f0; padding: 12px;">{action.get('echeance', 'Non définie')}</td>
                <td style="border: 1px solid #e2e8f0; padding: 12px;">{action.get('contexte', '')}</td>
            </tr>
            """
        
        html += "</tbody></table>"
        return html
    
    def _generate_risks_table_html(self, risks: list) -> str:
        """Generate HTML table for risks"""
        if not risks:
            return '<div style="color: #64748b; font-style: italic; text-align: center; padding: 20px;">Aucun risque identifié.</div>'
        
        html = """
        <table style="width: 100%; border-collapse: collapse; margin: 15px 0;">
            <thead>
                <tr><th style="border: 1px solid #e2e8f0; padding: 12px; text-align: left; background: #f8fafc; font-weight: bold; color: #4F46E5;">Risque</th><th style="border: 1px solid #e2e8f0; padding: 12px; text-align: left; background: #f8fafc; font-weight: bold; color: #4F46E5;">Catégorie</th><th style="border: 1px solid #e2e8f0; padding: 12px; text-align: left; background: #f8fafc; font-weight: bold; color: #4F46E5;">Probabilité</th><th style="border: 1px solid #e2e8f0; padding: 12px; text-align: left; background: #f8fafc; font-weight: bold; color: #4F46E5;">Impact</th><th style="border: 1px solid #e2e8f0; padding: 12px; text-align: left; background: #f8fafc; font-weight: bold; color: #4F46E5;">Mitigation</th><th style="border: 1px solid #e2e8f0; padding: 12px; text-align: left; background: #f8fafc; font-weight: bold; color: #4F46E5;">Responsable</th></tr>
            </thead>
            <tbody>
        """
        
        for risk in risks:
            prob = risk.get('probabilite', 'Moyenne')
            if prob == "Élevée":
                border_color = "#ef4444"
                bg_color = "#fef2f2"
            elif prob == "Moyenne":
                border_color = "#f59e0b"
                bg_color = "#fefbf2"
            else:
                border_color = "#10b981"
                bg_color = "#f0fdf4"
            
            html += f"""
            <tr style="border-left: 4px solid {border_color}; background: {bg_color};">
                <td style="border: 1px solid #e2e8f0; padding: 12px;">{risk.get('risque', '')}</td>
                <td style="border: 1px solid #e2e8f0; padding: 12px;">{risk.get('categorie', 'Général')}</td>
                <td style="border: 1px solid #e2e8f0; padding: 12px;">{risk.get('probabilite', 'Moyenne')}</td>
                <td style="border: 1px solid #e2e8f0; padding: 12px;">{risk.get('impact', '')}</td>
                <td style="border: 1px solid #e2e8f0; padding: 12px;">{risk.get('mitigations', risk.get('mitigation', ''))}</td>
                <td style="border: 1px solid #e2e8f0; padding: 12px;">{risk.get('responsableRisque', risk.get('responsable', ''))}</td>
            </tr>
            """
        
        html += "</tbody></table>"
        return html
    
    def _generate_chronological_html(self, chronological_data: list) -> str:
        """Generate HTML for chronological view"""
        html = '<div class="section" style="background: linear-gradient(135deg, #f093fb 0%, #f5576c 100%); color: white; padding: 20px; border-radius: 8px; margin-bottom: 25px;"><h2 style="color: white; border-left: 4px solid white; padding-left: 15px;">Vue chronologique de la réunion</h2>'
        
        if chronological_data:
            html += "<ul style='padding-left: 20px;'>"
            for event in chronological_data:
                if event and not event.startswith("/*"):
                    html += f"<li style='margin-bottom: 8px; line-height: 1.4;'>{event}</li>"
            html += "</ul>"
        else:
            html += '<div style="text-align: center; padding: 20px;">Aucune information chronologique.</div>'
        
        html += "</div>"
        return html
    
    def _generate_metrics_html(self, metrics: Dict[str, Any]) -> str:
        """Generate HTML for analysis metrics"""
        html = '<div class="section" style="background: linear-gradient(135deg, #4facfe 0%, #00f2fe 100%); color: white; padding: 20px; border-radius: 8px; margin-bottom: 25px;"><h2 style="color: white; border-left: 4px solid white; padding-left: 15px;">Métriques d\'analyse</h2>'
        
        metrics_list = [
            f"Segments analysés: {metrics.get('segmentsAnalyses', 0)}/{metrics.get('totalSegments', 0)}",
            f"Niveau de détail: {metrics.get('niveauDetaille', 'Non spécifié')}",
            f"Couverture des sujets: {metrics.get('couvertureSujets', 'Non spécifié')}",
            f"Qualité d'extraction: {metrics.get('qualiteExtraction', 'Non spécifié')}"
        ]
        
        html += "<ul style='padding-left: 20px;'>"
        for metric in metrics_list:
            html += f"<li style='margin-bottom: 8px; line-height: 1.4;'>{metric}</li>"
        html += "</ul>"
        
        html += "</div>"
        return html
    
    def _add_general_information(self, doc: Document, meta_data: Dict[str, Any]):
        """Add beautifully formatted general information section"""
        # Add title
        doc.add_heading("1. Informations générales", level=1)
        
        # Create a formatted table for meta information
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.allow_autofit = True
        
        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Élément'
        hdr_cells[1].text = 'Information'
        
        # Style header
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = self.brand_primary
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            # Add light background to header
            try:
                from docx.oxml.shared import qn
                tcPr = cell._tc.get_or_add_tcPr()
                shd = tcPr.find(qn('w:shd'))
                if shd is None:
                    from docx.oxml.parser import parse_xml
                    from docx.oxml.ns import nsdecls
                    shd = parse_xml(r'<w:shd {} w:fill="F1F5F9"/>'.format(nsdecls('w')))
                    tcPr.append(shd)
            except:
                pass  # Skip background color if it fails
        
        # Data rows
        info_items = [
            ("Projet", meta_data.get('projectName', 'Non spécifié')),
            ("Titre de la réunion", meta_data.get('meetingTitle', 'Non spécifié')),
            ("Type de réunion", meta_data.get('meetingType', 'Non spécifié')),
            ("Date de la réunion", self._format_date(meta_data.get('meetingDate', ''))),
            ("Durée", f"{meta_data.get('duration', 0)} minutes"),
            ("Participants attendus", str(meta_data.get('participantsExpected', 'Non spécifié'))),
            ("Participants détectés", ', '.join(meta_data.get('participantsDetected', []))),
        ]
        
        # Add instructions if present
        user_instructions = meta_data.get('userInstructions', '')
        if user_instructions:
            info_items.append(("Instructions spécifiques", user_instructions))
        
        # Add attendance analysis if present
        attendance_analysis = meta_data.get('attendanceAnalysis', '')
        if attendance_analysis:
            info_items.append(("Analyse de présence", attendance_analysis))
        
        for label, value in info_items:
            row_cells = table.add_row().cells
            row_cells[0].text = label
            row_cells[1].text = str(value)
            
            # Style label cell
            row_cells[0].paragraphs[0].runs[0].font.bold = True
            row_cells[0].paragraphs[0].runs[0].font.color.rgb = self.brand_secondary
            
            # Handle long text in value cell
            if len(str(value)) > 50:
                row_cells[1].paragraphs[0].runs[0].font.size = Pt(10)
        
        doc.add_paragraph()  # Add spacing
    
    def _format_date(self, date_str: str) -> str:
        """Format date string to be more readable"""
        if not date_str:
            return "Non spécifié"
        
        try:
            # Try to parse ISO format
            if 'T' in date_str:
                dt = datetime.fromisoformat(date_str.replace('Z', '+00:00'))
                return dt.strftime('%d/%m/%Y à %H:%M')
            else:
                return date_str
        except:
            return date_str
    
    def _add_chronological_view_word(self, doc: Document, chronological_data: list):
        """Add chronological timeline section to Word document"""
        if not chronological_data:
            return
            
        doc.add_heading("Vue chronologique de la réunion", level=1)
        
        for event in chronological_data:
            if event and not event.startswith("/*"):
                para = doc.add_paragraph()
                para.style = 'List Bullet'
                run = para.add_run(event)
                run.font.size = Pt(11)
        
        doc.add_paragraph()  # Add spacing
    
    def _add_analysis_metrics_word(self, doc: Document, metrics: Dict[str, Any]):
        """Add analysis quality metrics section to Word document"""
        if not metrics:
            return
            
        doc.add_heading("Métriques d'analyse", level=1)
        
        # Create a nice table for metrics
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.allow_autofit = True
        
        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Métrique'
        hdr_cells[1].text = 'Valeur'
        
        # Style header
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = self.brand_primary
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            # Add light background to header
            try:
                from docx.oxml.shared import qn
                tcPr = cell._tc.get_or_add_tcPr()
                shd = tcPr.find(qn('w:shd'))
                if shd is None:
                    from docx.oxml.parser import parse_xml
                    from docx.oxml.ns import nsdecls
                    shd = parse_xml(r'<w:shd {} w:fill="F1F5F9"/>'.format(nsdecls('w')))
                    tcPr.append(shd)
            except:
                pass  # Skip background color if it fails
        
        # Metrics data
        metrics_items = [
            ("Segments analysés", f"{metrics.get('segmentsAnalyses', 0)}/{metrics.get('totalSegments', 0)}"),
            ("Niveau de détail", metrics.get('niveauDetaille', 'Non spécifié')),
            ("Couverture des sujets", metrics.get('couvertureSujets', 'Non spécifié')),
            ("Qualité d'extraction", metrics.get('qualiteExtraction', 'Non spécifié'))
        ]
        
        for label, value in metrics_items:
            row_cells = table.add_row().cells
            row_cells[0].text = label
            row_cells[1].text = str(value)
            
            # Style label cell
            row_cells[0].paragraphs[0].runs[0].font.bold = True
            row_cells[0].paragraphs[0].runs[0].font.color.rgb = self.brand_secondary
        
        doc.add_paragraph()  # Add spacing
